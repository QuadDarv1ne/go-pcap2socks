from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_paragraph_font(paragraph, font_name, size, bold=False):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        # Для кириллицы обязательно менять шрифт и в eastern europe
        run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font_name)

def add_custom_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'Times New Roman')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_custom_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'Times New Roman')
    p.paragraph_format.space_after = Pt(6)
    # Выравнивание по ширине
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    # Clear default runs and add formatted one
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'Times New Roman')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_code_block(doc, code, font_size=10):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(font_size)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

doc = Document()

# Убираем отступы по умолчанию
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

add_custom_heading(doc, "Анализ и пути улучшения проекта go-pcap2socks")
add_custom_text(doc, "Проект go-pcap2socks представляет собой современное инженерное решение для перехвата сетевого трафика на уровне gVisor userspace stack и его проброса через SOCKS5-прокси. Данный документ содержит подробный анализ текущей архитектуры и предлагает комплексные идеи по улучшению проекта для перевода его из статуса «рабочий концепт» в статус надежного production-инструмента.")

add_custom_heading(doc, "1. Текущая архитектура проекта")
add_custom_text(doc, "Проект использует передовой подход с применением gVisor — userspace network stack от Google, который предоставляет полную реализацию TCP/IP стека в пользовательском пространстве. Это eliminates необходимость в raw sockets и ручном управлении TCP состояниями.")
add_bullet(doc, "core/stack.go — создание и настройка gVisor stack с поддержкой IPv4/IPv6, TCP, UDP, ICMP")
add_bullet(doc, "core/tcp.go — TCP forwarder на базе gVisor, обработка входящих соединений")
add_bullet(doc, "core/udpforwarder.go — UDP NAT с использованием sing-box udpnat.Service")
add_bullet(doc, "proxy/socks5.go — реализация SOCKS5 клиента с connection pooling")
add_bullet(doc, "dns/ — модуль для DNS resolution через DoH/DoT")

add_custom_heading(doc, "2. Критические аспекты стабильности")
add_custom_text(doc, "Несмотря на использование gVisor, остаются важные вопросы управления жизненным циклом соединений и обработки ошибок.")
add_bullet(doc, "Утечки Goroutine и управление состоянием: При обрыве соединения (получение пакетов FIN/RST от gVisor) или ошибке SOCKS5-сервера необходимо жестко закрывать все горутины. Улучшение: внедрить context.WithCancel на каждое UDP/TCP соединение и вызывать cancel() в defer-блоках.")
add_bullet(doc, "Обработка TCP State Machine в gVisor: gVisor автоматически управляет TCP состояниями (handshake, retransmission, flow control), но приложение должно корректно обрабатывать события закрытия. Стратегическое улучшение: использовать gVisor Endpoint State notifications для детектирования закрытия соединений.")
add_bullet(doc, "Контроль памяти в UDP NAT: Текущая реализация udpforwarder.go использует атомарные счетчики сессий, но нет жесткого лимита. Улучшение: добавить maxUDPSessions (уже есть константа, но нужна динамическая настройка) и dropped packet metrics.")

add_custom_heading(doc, "3. Оптимизация производительности")
add_custom_text(doc, "gVisor значительно снижает overhead по сравнению с raw sockets, но правильная настройка concurrency и буферов критически важна.")
add_bullet(doc, "Thread-safe ConnTrack Map: Хранение состояний активных соединений в стандартной map с sync.RWMutex создает узкие места. Улучшение: использовать sync.Map для hot path или библиотеку github.com/orcaman/concurrent-map с шардированием.")
add_bullet(doc, "Оптимизация буферов Relay: В proxy/socks5.go используются буферы 32 КБ (pool.Get(32*1024)), что оптимально для TCP. Для UDP рекомендуется 2-4 КБ из-за MTU ограничений.")
add_bullet(doc, "Connection Pooling для SOCKS5: В Socks5 struct уже есть connPool, но можно улучшить: добавить pre-warming пула при старте, health check фоновую горутину для проверки доступности прокси.")

add_custom_heading(doc, "4. Новые функциональные возможности")
add_custom_text(doc, "Для практического применения в production не хватает нескольких ключевых функций.")
add_bullet(doc, "DNS Hijacking (Фейковый DNS): Перехватывать DNS запросы (порт 53) и возвращать фиктивный IP из диапазона 198.51.100.0/24. Когда приложение подключится на этот IP, прокси поймет соответствие и подключится к SOCKS5 по домену.")
add_bullet(doc, "Аутентификация Username/Password: В proxy/socks5.go уже есть поддержка user/pass, но нужна валидация и обработка ошибок auth.")
add_bullet(doc, "Маршрутизация и White/Black Lists: Возможность исключать локальные подсети (192.168.0.0/16, 10.0.0.0/8) из проброса.")
add_bullet(doc, "UPnP/NAT-PMP: В проекте уже есть upnp модуль, но нужна интеграция с основным flow для автоматического проброса портов.")
add_bullet(doc, "WireGuard интеграция: В проекте есть wireguard модуль, но нужна документация по использованию.")

add_custom_heading(doc, "5. Developer Experience и мониторинг")
add_custom_text(doc, "Удобство отладки и наблюдения за состоянием программы критически важно.")
add_bullet(doc, "Структурированное логирование: Проект уже использует log/slog — отлично! Добавить контекст: src IP, dst IP:port, connection ID.")
add_bullet(doc, "Экспорт метрик (Prometheus): Добавить HTTP-сервер на :9090/metrics. Ключевые метрики: active_tcp_sessions, active_udp_sessions, socks5_pool_hits, socks5_pool_misses, udp_dropped_packets, bytes_proxied_total.")
add_bullet(doc, "Конфигурационный файл: Поддержка YAML/TOML через Viper. Избегать перегрузки CLI аргументами.")
add_bullet(doc, "Health Checks для SOCKS5: В Socks5 struct уже есть CheckHealth(), но нужен фоновый worker с интервалом 30-60 сек.")

add_custom_heading(doc, "6. Безопасность")
add_custom_text(doc, "Специфичные требования к безопасности для сетевого прокси.")
add_bullet(doc, "Сброс привилегий (Drop Privileges): gVisor требует минимальных привилегий. После создания stack, сбросить capabilities до минимума.")
add_bullet(doc, "Защита от SSRF: Запретить проброс трафика к внутренним интерфейсам (127.0.0.0/8, 169.254.0.0/16) через SOCKS.")
add_bullet(doc, "Rate Limiting: В проекте есть ratelimit модуль — использовать для ограничения новых соединений в секунду.")

add_custom_heading(doc, "7. Приоритетный план развития (Roadmap)")
add_custom_text(doc, "Рекомендуемый порядок внедрения изменений:")
add_bullet(doc, "Шаг 1: Рефакторинг concurrency. Context management для UDP/TCP сессий, graceful shutdown.")
add_bullet(doc, "Шаг 2: Интеграция логгера slog с контекстом и добавление базовых метрик (active sessions, pool stats).")
add_bullet(doc, "Шаг 3: DNS Hijacking модуль для повседневного использования.")
add_bullet(doc, "Шаг 4: Health check worker для SOCKS5 пула с автоматическим исключением unhealthy прокси.")
add_bullet(doc, "Шаг 5: Prometheus metrics exporter с интеграцией в observability модуль.")

doc.save('go-pcap2socks_improvements.docx')
print("Документ успешно создан: go-pcap2socks_improvements.docx")
